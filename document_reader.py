"""Validated document-reading boundary for downstream conversion and LLM work."""

from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from docx import Document


@dataclass(frozen=True)
class SourceDocument:
    """A validated source document and the reader required to handle it."""

    path: Path
    kind: Literal["docx", "text"]
    content: Document | str


class DocumentReader:
    """Open supported source documents without exposing file handling to the LLM layer."""

    TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".rst", ".csv", ".json", ".yaml", ".yml", ".log"}

    def read(self, document_path: str) -> SourceDocument:
        path = Path(document_path).expanduser().resolve()
        if not path.is_file():
            raise FileNotFoundError(f"Document not found: {path}")
        if path.suffix.lower() == ".docx":
            return SourceDocument(path=path, kind="docx", content=Document(path))
        if path.suffix.lower() in self.TEXT_EXTENSIONS:
            return SourceDocument(path=path, kind="text", content=path.read_text(encoding="utf-8-sig"))
        supported = ", ".join([".docx", *sorted(self.TEXT_EXTENSIONS)])
        raise ValueError(f"Unsupported document type '{path.suffix or '(no extension)'}'. Supported types: {supported}.")

    def read_docx(self, document_path: str) -> Document:
        source = self.read(document_path)
        if source.kind != "docx":
            raise ValueError("This operation requires a .docx document.")
        return source.content
